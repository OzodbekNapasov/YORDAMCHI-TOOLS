import os
import re
import copy
import json
import zipfile
import io
from datetime import datetime
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

from docbot_config import TEMPLATES as DOCBOT_TEMPLATES, find_template_file

DISTRICT_DOCTORS = {
    "Shahrisabz shahar": "O.Norboyev",
    "Kitob tuman": "A.Hasanov",
    "Yakkabog' tuman": "S.B.Jo’rayev",
    "Shahrisabz tuman": "Z.Esanov",
    "Chiroqchi tuman": "Sh.Ro'ziqulov",
    "Qamashi tuman": "Avazov Shuxrat Shukullayevich"
}

STANDARD_DISTRICTS = [
    "Shahrisabz shahar",
    "Kitob tuman",
    "Yakkabog' tuman",
    "Shahrisabz tuman",
    "Chiroqchi tuman",
    "Qamashi tuman"
]


def format_amaliyot_muddati(start_date_str: str, end_date_str: str) -> str:
    """Sanani '2026-yil 08-iyunidan  2026-yil 06-iyuligacha' formatiga o'tkazish"""
    uzbek_months = {
        1: ("yanvar", "yanvaridan", "yanvarigacha"),
        2: ("fevral", "fevralidan", "fevraligacha"),
        3: ("mart", "martidan", "martigacha"),
        4: ("aprel", "aprelidan", "apreligacha"),
        5: ("may", "mayidan", "mayigacha"),
        6: ("iyun", "iyunidan", "iyunigacha"),
        7: ("iyul", "iyulidan", "iyuligacha"),
        8: ("avgust", "avgustidan", "avgustigacha"),
        9: ("sentabr", "sentabridan", "sentabrigacha"),
        10: ("oktabr", "oktabridan", "oktabrigacha"),
        11: ("noyabr", "noyabridan", "noyabrigacha"),
        12: ("dekabr", "dekabridan", "dekabrigacha")
    }

    try:
        s_dt = datetime.strptime(start_date_str.strip(), "%d.%m.%Y")
        e_dt = datetime.strptime(end_date_str.strip(), "%d.%m.%Y")

        s_m_str = uzbek_months[s_dt.month][1]  # -dan
        e_m_str = uzbek_months[e_dt.month][2]  # -gacha

        s_day = f"{s_dt.day:02d}"
        e_day = f"{e_dt.day:02d}"

        return f"{s_dt.year}-yil {s_day}-{s_m_str}  {e_dt.year}-yil {e_day}-{e_m_str}"
    except Exception:
        return f"{start_date_str} dan {end_date_str} gacha"


def set_cell_formatted(cell, text, font_name="Times New Roman", font_size=10.5, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, width=None):
    """Jadval katakchasiga Times New Roman shriftida ixcham va tekis matn joylash"""
    cell.text = ""
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0.5)
    p.paragraph_format.space_after = Pt(0.5)
    p.paragraph_format.line_spacing = 1.0

    run = p.add_run(str(text))
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)

    # Windows va Word uchun Times New Roman rFonts xususiyatini o'rnatish
    rPr = run._r.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}" w:cs="{font_name}"/>')
    rPr.append(rFonts)

    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if width:
        cell.width = width


def fill_amaliyot_template(template_path: str, data: dict, output_path: str):
    """
    Amaliyot shabloni (.docx) ni to'liq to'ldirib, talabalar jadvalini dinamik kengaytiradi.
    Times New Roman shrifti, qat'iy ustun kengliklari va ixcham qatorlar bilan formatlaydi.
    """
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Amaliyot shabloni topilmadi: {template_path}")

    doc = docx.Document(template_path)

    buyruq_raqami = data.get("buyruq_raqami", "").strip() or "____"
    buyruq_sanasi = data.get("buyruq_sanasi", "").strip() or datetime.now().strftime("%d.%m.%Y")
    tumani = data.get("tumani", "").strip() or "Shahrisabz shahar"
    shu_tuman_shifokori = data.get("shu_tuman_shifokori", "").strip() or DISTRICT_DOCTORS.get(tumani, "Bosh shifokor")
    oquv_yili = data.get("oquv_yili", "").strip() or "2025/2026"
    kursi = str(data.get("kursi", "1")).strip()
    
    # Guruhlar ro'yxatini shakllantirish
    raw_guruhlar = data.get("guruhlar", [])
    if isinstance(raw_guruhlar, str):
        guruhlar = [g.strip() for g in re.split(r'[,; ]+', raw_guruhlar) if g.strip()]
    elif isinstance(raw_guruhlar, list):
        guruhlar = [str(g).strip() for g in raw_guruhlar if str(g).strip()]
    else:
        guruhlar = []

    guruhlar_str = ", ".join(guruhlar) if guruhlar else "101, 102"

    start_date = data.get("start_date", "").strip() or "08.06.2026"
    end_date = data.get("end_date", "").strip() or "06.07.2026"

    amaliyot_muddati = data.get("amaliyot_muddati", "").strip()
    if not amaliyot_muddati and start_date and end_date:
        amaliyot_muddati = format_amaliyot_muddati(start_date, end_date)
    if not amaliyot_muddati:
        amaliyot_muddati = "2026-yil 08-iyunidan  2026-yil 06-iyuligacha"

    replacements = {
        "{{buyruq_raqami}}": buyruq_raqami,
        "{{buyruq_sanasi}}": buyruq_sanasi,
        "{{tumani}}": tumani,
        "{{shu_tuman_shifokori}}": shu_tuman_shifokori,
        "{{oquv_yili}}": oquv_yili,
        "{{kursi}}": kursi,
        "{{guruh_1}}, {{guruh_2}}, {{guruh_3}}": guruhlar_str,
        "{{guruh_1}},{{guruh_2}},{{guruh_3}}": guruhlar_str,
        "{{guruh_1}}": guruhlar[0] if len(guruhlar) > 0 else "",
        "{{guruh_2}}": guruhlar[1] if len(guruhlar) > 1 else "",
        "{{guruh_3}}": guruhlar[2] if len(guruhlar) > 2 else "",
        "{{amaliyot_muddati}}": amaliyot_muddati,
        "{{amaliyot_boshlanish_sanasi}}": start_date,
        "{{amaliyot_tugash_sanasi}}": end_date
    }

    def _replace_in_p(p):
        full_text = p.text
        has_match = False
        for k, v in replacements.items():
            if k in full_text:
                full_text = full_text.replace(k, v)
                has_match = True
        if has_match:
            if p.runs:
                p.runs[0].text = full_text
                for r in p.runs[1:]:
                    r.text = ""
            else:
                p.text = full_text

    # Paragraf matnlarini almashtirish
    for p in doc.paragraphs:
        _replace_in_p(p)

    # 1 va 2-jadvallardagi matnlarni almashtirish
    for t_idx, table in enumerate(doc.tables):
        if t_idx < 2:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        _replace_in_p(p)

    # 3-Jadval: Talabalar ro'yxati jadvali (Table 2)
    students = data.get("students", [])
    if len(doc.tables) >= 3 and students:
        t = doc.tables[2]

        col_widths = [
            Inches(0.42),  # 0: T/r
            Inches(0.72),  # 1: Guruhi
            Inches(2.55),  # 2: F.I.SH (Keng va 1 qatorga sig'adigan)
            Inches(0.92),  # 3: Boshlanishi
            Inches(0.92),  # 4: Tugashi
            Inches(0.68),  # 5: Bahosi
            Inches(0.68)   # 6: Imzo
        ]

        # Keraksiz shablon qatorlarini tozalash
        while len(t.rows) > len(students) + 1:
            tr = t.rows[-1]._tr
            t._tbl.remove(tr)

        # Agar talabalar ko'proq bo'lsa yangi qatorlar qo'shish
        while len(t.rows) < len(students) + 1:
            new_tr = copy.deepcopy(t.rows[1]._tr)
            t._tbl.append(new_tr)

        # 1. Sarlavha qatorini (Row 0) chiroyli Times New Roman Bold formatlash
        header_titles = [
            "T/r", "Guruhi", "F.I.SH", 
            "Amaliyot boshlanishi vaqti", "Amaliyot tugash vaqti", 
            "Amaliyot bahosi", "Rahbar imzosi"
        ]
        if len(t.rows) > 0:
            for c_idx, cell in enumerate(t.rows[0].cells):
                title = header_titles[c_idx] if c_idx < len(header_titles) else cell.text
                set_cell_formatted(
                    cell, title, 
                    font_name="Times New Roman", font_size=10.0, bold=True, 
                    align=WD_ALIGN_PARAGRAPH.CENTER, width=col_widths[c_idx] if c_idx < len(col_widths) else None
                )

        # 2. Talabalar qatorlarini (Row 1..N) to'ldirish va Times New Roman formatlash
        for idx, st in enumerate(students):
            row = t.rows[idx + 1]
            st_fio = st.get("fio", "").strip()
            st_guruh = st.get("guruhi", "").strip() or (guruhlar[0] if guruhlar else "")
            st_start = st.get("start_date", "").strip() or start_date
            st_end = st.get("end_date", "").strip() or end_date

            fio_font_size = 9.5 if len(st_fio) > 30 else 10.5

            if len(row.cells) > 0:
                set_cell_formatted(row.cells[0], f"{idx + 1}.", "Times New Roman", 10.5, False, WD_ALIGN_PARAGRAPH.CENTER, col_widths[0])
            if len(row.cells) > 1:
                set_cell_formatted(row.cells[1], st_guruh, "Times New Roman", 10.5, False, WD_ALIGN_PARAGRAPH.CENTER, col_widths[1])
            if len(row.cells) > 2:
                set_cell_formatted(row.cells[2], st_fio, "Times New Roman", fio_font_size, False, WD_ALIGN_PARAGRAPH.LEFT, col_widths[2])
            if len(row.cells) > 3:
                set_cell_formatted(row.cells[3], st_start, "Times New Roman", 10.5, False, WD_ALIGN_PARAGRAPH.CENTER, col_widths[3])
            if len(row.cells) > 4:
                set_cell_formatted(row.cells[4], st_end, "Times New Roman", 10.5, False, WD_ALIGN_PARAGRAPH.CENTER, col_widths[4])
            if len(row.cells) > 5:
                set_cell_formatted(row.cells[5], "", "Times New Roman", 10.5, False, WD_ALIGN_PARAGRAPH.CENTER, col_widths[5])
            if len(row.cells) > 6:
                set_cell_formatted(row.cells[6], "", "Times New Roman", 10.5, False, WD_ALIGN_PARAGRAPH.CENTER, col_widths[6])

        for r in t.rows:
            for c_i, c in enumerate(r.cells):
                if c_i < len(col_widths):
                    c.width = col_widths[c_i]

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path


def generate_sample_survey_excel(output_path: str = None) -> bytes:
    """
    So'rovnoma uchun zamonaviy va chiroyli dizayndagi namuna Excel (.xlsx) faylini yaratadi.
    Tumanlar, guruhlar, talabalar va sanalar bilan to'ldirilgan namunaviy qatorlarni o'z ichiga oladi.
    """
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Amaliyot So'rovnomasi"

    ws.views.sheetView[0].showGridLines = True

    HEADER_FILL = PatternFill(start_color="0D5C56", end_color="0D5C56", fill_type="solid")
    ALT_ROW_FILL = PatternFill(start_color="F7FAFC", end_color="F7FAFC", fill_type="solid")

    FONT_HEADER = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
    FONT_TITLE = Font(name="Calibri", size=14, bold=True, color="0D5C56")
    FONT_NOTE = Font(name="Calibri", size=10, italic=True, color="4A5568")
    FONT_DATA = Font(name="Calibri", size=10.5, bold=False, color="1A202C")
    FONT_DATA_BOLD = Font(name="Calibri", size=10.5, bold=True, color="1A202C")

    THIN_BORDER = Border(
        left=Side(style="thin", color="CBD5E0"),
        right=Side(style="thin", color="CBD5E0"),
        top=Side(style="thin", color="CBD5E0"),
        bottom=Side(style="thin", color="CBD5E0")
    )
    HEADER_BORDER = Border(
        left=Side(style="thin", color="0D5C56"),
        right=Side(style="thin", color="0D5C56"),
        top=Side(style="thin", color="0D5C56"),
        bottom=Side(style="medium", color="0B3B39")
    )

    ws.merge_cells("A1:H1")
    title_cell = ws["A1"]
    title_cell.value = "MALAKAVIY AMALIYOT TALABALAR SO'ROVNOMASI VA TUMANLAR TAQSIMOTI"
    title_cell.font = FONT_TITLE
    title_cell.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 26

    ws.merge_cells("A2:H2")
    note_cell = ws["A2"]
    note_cell.value = "Eslatma: 'Amaliyot Tumani' ustuniga tuman nomini aniq yozing (Shahrisabz shahar, Kitob tuman, Yakkabog' tuman, Chiroqchi tuman, Qamashi tuman, Shahrisabz tuman va h.k.)."
    note_cell.font = FONT_NOTE
    note_cell.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[2].height = 20

    headers = [
        "T/r",
        "Guruhi",
        "Talabaning F.I.SH",
        "Amaliyot Tumani",
        "Boshlanish sanasi",
        "Tugash sanasi",
        "Telefon raqami",
        "Amaliyot muassasasi / Shifoxona"
    ]

    ws.row_dimensions[4].height = 28
    for col_num, header_title in enumerate(headers, 1):
        cell = ws.cell(row=4, column=col_num)
        cell.value = header_title
        cell.font = FONT_HEADER
        cell.fill = HEADER_FILL
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = HEADER_BORDER

    sample_rows = [
        (1, "201", "Rahmatova Shaxnoza Sherzod qizi", "Shahrisabz shahar", "08.06.2026", "06.07.2026", "+998901234567", "Shahrisabz ShTTB Markaziy Shifoxonasi"),
        (2, "201", "Botirova Gulbahor Olimovna", "Shahrisabz shahar", "08.06.2026", "06.07.2026", "+998912345678", "Shahrisabz ShTTB 1-sonli Poliklinika"),
        (3, "202", "Asraliyev Asilbek Bekmurod o'g'li", "Kitob tuman", "08.06.2026", "06.07.2026", "+998933456789", "Kitob TTB Markaziy Shifoxonasi"),
        (4, "202", "Meyliyev Ruslan Rustam o'g'li", "Kitob tuman", "08.06.2026", "06.07.2026", "+998974567890", "Kitob TTB Shoshilinch Bo'limi"),
        (5, "203", "Nazarova Dilnoza Farxod qizi", "Yakkabog' tuman", "08.06.2026", "06.07.2026", "+998995678901", "Yakkabog' TTB Markaziy Poliklinikasi"),
        (6, "203", "Qodirov Jasur Anvar o'g'li", "Yakkabog' tuman", "08.06.2026", "06.07.2026", "+998906789012", "Yakkabog' TTB 2-sonli Shifoxonasi"),
        (7, "204", "Eshmurodov Bobur Shavkat o'g'li", "Chiroqchi tuman", "08.06.2026", "06.07.2026", "+998917890123", "Chiroqchi TTB Markaziy Shifoxonasi"),
        (8, "204", "Xoliqova Madina Zafar qizi", "Qamashi tuman", "08.06.2026", "06.07.2026", "+998988901234", "Qamashi TTB Markaziy Shifoxonasi"),
        (9, "204", "Jumanov Sardor Bahodir o'g'li", "Shahrisabz tuman", "08.06.2026", "06.07.2026", "+998909012345", "Shahrisabz Tuman TTB Shifoxonasi")
    ]

    for r_idx, row_data in enumerate(sample_rows, 5):
        ws.row_dimensions[r_idx].height = 22
        is_alt = (r_idx % 2 == 0)
        row_fill = ALT_ROW_FILL if is_alt else None

        for c_idx, val in enumerate(row_data, 1):
            cell = ws.cell(row=r_idx, column=c_idx)
            cell.value = val
            cell.font = FONT_DATA_BOLD if c_idx in [1, 2, 4] else FONT_DATA
            if row_fill:
                cell.fill = row_fill
            cell.border = THIN_BORDER

            if c_idx in [1, 2, 5, 6, 7]:
                cell.alignment = Alignment(horizontal="center", vertical="center")
            elif c_idx == 4:
                cell.alignment = Alignment(horizontal="center", vertical="center")
            else:
                cell.alignment = Alignment(horizontal="left", vertical="center")

    col_widths = {
        "A": 7,
        "B": 12,
        "C": 36,
        "D": 22,
        "E": 18,
        "F": 18,
        "G": 18,
        "H": 36
    }
    for col_letter, width in col_widths.items():
        ws.column_dimensions[col_letter].width = width

    if output_path:
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        wb.save(output_path)
        return output_path
    else:
        stream = io.BytesIO()
        wb.save(stream)
        stream.seek(0)
        return stream.getvalue()


def parse_survey_excel(file_bytes_or_path, default_start="08.06.2026", default_end="06.07.2026") -> list:
    """
    Yuklangan Excel (.xlsx) faylini o'qib, talabalar so'rovnomasi ro'yxatini qaytaradi.
    Moslashuvchan: ustun sarlavhalarini aqlli aniqlaydi.
    """
    if isinstance(file_bytes_or_path, (bytes, bytearray)):
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes_or_path), data_only=True)
    elif hasattr(file_bytes_or_path, 'read'):
        wb = openpyxl.load_workbook(file_bytes_or_path, data_only=True)
    else:
        wb = openpyxl.load_workbook(file_bytes_or_path, data_only=True)

    ws = wb.active

    best_row_idx = None
    best_score = 0
    col_map = {
        "tr": None,
        "guruhi": None,
        "fio": None,
        "tumani": None,
        "start_date": None,
        "end_date": None,
        "phone": None,
        "organization": None
    }

    for r_idx in range(1, min(15, ws.max_row + 1)):
        row_vals = [str(ws.cell(row=r_idx, column=c).value or "").strip().lower() for c in range(1, ws.max_column + 1)]
        cur_map = {}
        score = 0
        for c_idx, v in enumerate(row_vals, 1):
            if not v:
                continue
            if "t/r" in v or "№" in v or "nomer" in v or v == "tr":
                cur_map["tr"] = c_idx
                score += 2
            elif "guruh" in v:
                cur_map["guruhi"] = c_idx
                score += 3
            elif "f.i.sh" in v or "fio" in v or "familiya" in v or "talabaning" in v or ("talaba" in v and "ism" in v):
                cur_map["fio"] = c_idx
                score += 4
            elif "tuman" in v or "shahar" in v or "amaliyot joy" in v or "amaliyot tumani" in v:
                cur_map["tumani"] = c_idx
                score += 3
            elif "boshlanish" in v or "boshla" in v:
                cur_map["start_date"] = c_idx
                score += 2
            elif "tugash" in v or "tuga" in v:
                cur_map["end_date"] = c_idx
                score += 2
            elif "tel" in v or "telefon" in v or "aloqa" in v:
                cur_map["phone"] = c_idx
                score += 1
            elif "muassasa" in v or "shifoxona" in v or "baza" in v or "poliklinika" in v:
                cur_map["organization"] = c_idx
                score += 1

        if score > best_score and ("fio" in cur_map or "guruhi" in cur_map or "tumani" in cur_map):
            best_score = score
            best_row_idx = r_idx
            for k, val in cur_map.items():
                col_map[k] = val

    if best_row_idx is None:
        # Sarlavha topilmadi — ustunlar soniga qarab avtomatik moslash
        max_c = min(ws.max_column, 8)
        if max_c <= 2:
            col_map = {
                "tr": None,
                "guruhi": 1,
                "fio": 2,
                "tumani": None,
                "start_date": None,
                "end_date": None,
                "phone": None,
                "organization": None
            }
        elif max_c == 3:
            # 1-ustun raqam (T/r) bo'lsa: 1=tr, 2=guruhi, 3=fio
            first_val = str(ws.cell(row=1, column=1).value or "").strip()
            if first_val.isdigit() and len(first_val) <= 4:
                col_map = {"tr": 1, "guruhi": 2, "fio": 3, "tumani": None, "start_date": None, "end_date": None, "phone": None, "organization": None}
            else:
                col_map = {"tr": None, "guruhi": 1, "fio": 2, "tumani": 3, "start_date": None, "end_date": None, "phone": None, "organization": None}
        else:
            col_map = {
                "tr": 1,
                "guruhi": 2,
                "fio": 3,
                "tumani": 4,
                "start_date": 5,
                "end_date": 6,
                "phone": 7,
                "organization": 8
            }

        # 1-qator sarlavhami yoki ma'lumot ekanini tekshirish
        r1_c1 = str(ws.cell(row=1, column=1).value or "").strip().lower()
        r1_c2 = str(ws.cell(row=1, column=2).value or "").strip().lower()
        is_r1_header = any(k in r1_c1 or k in r1_c2 for k in ["guruh", "fio", "f.i.sh", "ism", "familiya", "t/r", "№", "name"])
        start_reading_row = 2 if is_r1_header else 1
    else:
        header_row_idx = best_row_idx
        if not col_map["fio"]:
            col_map["fio"] = 2 if ws.max_column <= 2 else 3
        if not col_map["guruhi"]:
            col_map["guruhi"] = 1 if ws.max_column <= 2 else 2
        start_reading_row = header_row_idx + 1

    records = []
    for r_idx in range(start_reading_row, ws.max_row + 1):
        fio_val = str(ws.cell(row=r_idx, column=col_map["fio"]).value or "").strip() if col_map["fio"] else ""
        if not fio_val or fio_val.lower() in ["none", "null", "f.i.sh", "talabaning f.i.sh"]:
            continue

        guruh_val = str(ws.cell(row=r_idx, column=col_map["guruhi"]).value or "").strip() if col_map["guruhi"] else ""
        if guruh_val.lower() in ["none", "null"]:
            guruh_val = ""
        if guruh_val.endswith(".0"):
            guruh_val = guruh_val[:-2]

        tuman_val = str(ws.cell(row=r_idx, column=col_map["tumani"]).value or "").strip() if col_map["tumani"] else ""
        if not tuman_val or tuman_val.lower() in ["none", "null"]:
            tuman_val = "Shahrisabz shahar"

        s_date_val = str(ws.cell(row=r_idx, column=col_map["start_date"]).value or "").strip() if col_map["start_date"] else ""
        if not s_date_val or s_date_val.lower() in ["none", "null"]:
            s_date_val = default_start
        else:
            if isinstance(ws.cell(row=r_idx, column=col_map["start_date"]).value, datetime):
                s_date_val = ws.cell(row=r_idx, column=col_map["start_date"]).value.strftime("%d.%m.%Y")

        e_date_val = str(ws.cell(row=r_idx, column=col_map["end_date"]).value or "").strip() if col_map["end_date"] else ""
        if not e_date_val or e_date_val.lower() in ["none", "null"]:
            e_date_val = default_end
        else:
            if isinstance(ws.cell(row=r_idx, column=col_map["end_date"]).value, datetime):
                e_date_val = ws.cell(row=r_idx, column=col_map["end_date"]).value.strftime("%d.%m.%Y")

        phone_val = str(ws.cell(row=r_idx, column=col_map["phone"]).value or "").strip() if col_map["phone"] else ""
        if phone_val.lower() in ["none", "null"]:
            phone_val = ""

        org_val = str(ws.cell(row=r_idx, column=col_map["organization"]).value or "").strip() if col_map["organization"] else ""
        if org_val.lower() in ["none", "null"]:
            org_val = ""

        records.append({
            "guruhi": guruh_val,
            "fio": fio_val,
            "tumani": tuman_val,
            "start_date": s_date_val,
            "end_date": e_date_val,
            "phone": phone_val,
            "organization": org_val
        })

    return records


def generate_all_district_orders(template_path: str, semester_data: dict, survey_students: list, output_dir: str) -> dict:
    """
    So'rovnomadagi barcha talabalarni tumanlar bo'yicha ajratib, har bir tuman uchun alohida Word (.docx)
    buyrug'ini yaratadi va barcha fayllarni bitta ZIP paketga arxivlaydi.
    """
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Amaliyot shabloni topilmadi: {template_path}")

    os.makedirs(output_dir, exist_ok=True)

    district_groups = {}
    for st in survey_students:
        tum = st.get("tumani", "").strip() or "Shahrisabz shahar"
        if tum not in district_groups:
            district_groups[tum] = []
        district_groups[tum].append(st)

    generated_files = []
    zip_filename = f"Amaliyot_Barcha_Tumanlar_Buyruqlari_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
    zip_path = os.path.join(output_dir, zip_filename)

    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for tumani, students in district_groups.items():
            doctor = DISTRICT_DOCTORS.get(tumani, semester_data.get("shu_tuman_shifokori", "Bosh shifokor"))
            
            dist_groups = sorted(list(set(s.get("guruhi", "").strip() for s in students if s.get("guruhi", "").strip())))
            dist_groups_str = ", ".join(dist_groups) if dist_groups else semester_data.get("guruhlar_str", "Guruh")

            start_date = semester_data.get("start_date", "08.06.2026")
            end_date = semester_data.get("end_date", "06.07.2026")
            amaliyot_muddati = semester_data.get("amaliyot_muddati") or format_amaliyot_muddati(start_date, end_date)

            order_data = {
                "buyruq_raqami": semester_data.get("buyruq_raqami", "____"),
                "buyruq_sanasi": semester_data.get("buyruq_sanasi", datetime.now().strftime("%d.%m.%Y")),
                "tumani": tumani,
                "shu_tuman_shifokori": doctor,
                "oquv_yili": semester_data.get("oquv_yili", "2025/2026"),
                "kursi": semester_data.get("kursi", "1"),
                "guruhlar": dist_groups or semester_data.get("guruhlar", []),
                "amaliyot_muddati": amaliyot_muddati,
                "start_date": start_date,
                "end_date": end_date,
                "students": students
            }

            clean_tumani = re.sub(r'[\\/*?:"<>|]', "", tumani).strip()
            clean_grp = re.sub(r'[\\/*?:"<>|]', "", dist_groups_str).strip()
            docx_filename = f"{clean_tumani} - {clean_grp} - {len(students)} ta talaba.docx"
            docx_file_path = os.path.join(output_dir, docx_filename)

            fill_amaliyot_template(template_path, order_data, docx_file_path)

            zipf.write(docx_file_path, arcname=docx_filename)

            generated_files.append({
                "tumani": tumani,
                "shifokor": doctor,
                "guruhlar": dist_groups_str,
                "students_count": len(students),
                "filename": docx_filename,
                "file_path": docx_file_path
            })

    return {
        "success": True,
        "zip_path": zip_path,
        "zip_filename": zip_filename,
        "total_districts": len(district_groups),
        "total_students": len(survey_students),
        "files": generated_files
    }
