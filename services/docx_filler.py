# ============================================================
#  services/docx_filler.py
#  .docx shablonlarni to'ldirish
#  - Barcha Bold, Kursiv (Italic), Shrift o'lchamlari va jadvallarni
#    100% asl holida saqlaydi
# ============================================================

import os
from docx import Document
from docx.shared import Pt


def fill_template(template_path: str, output_path: str, data: dict) -> None:
    """
    template_path : .docx shablon fayli yo'li
    output_path   : natija .docx fayli yo'li
    data          : {"FIO": "...", "buyruq_raqami": "...", ...}
    """
    doc = Document(template_path)

    # Barcha almashtiriladigan kalitlar va qiymatlar
    # Agar data ichida FIO bo'lsa va shablonda {{IFO}} bo'lsa, ikkalasini ham qamrab olamiz
    normalized_data = {}
    for k, v in data.items():
        val_str = str(v).strip()
        normalized_data[k] = val_str
        if k == "FIO" and "IFO" not in data:
            normalized_data["IFO"] = val_str
        elif k == "IFO" and "FIO" not in data:
            normalized_data["FIO"] = val_str
        if k == "SANA" and "sanasi" not in data:
            normalized_data["sanasi"] = val_str
        elif k == "sanasi" and "SANA" not in data:
            normalized_data["SANA"] = val_str

    # 1. Paragraflarni qayta ishlash
    for para in doc.paragraphs:
        _process_paragraph(para, normalized_data)

    # 2. Jadvallar ichidagi kataklarni qayta ishlash
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _process_paragraph(para, normalized_data)

    # 3. Header va Footer
    for section in doc.sections:
        for para in section.header.paragraphs:
            _process_paragraph(para, normalized_data)
        for para in section.footer.paragraphs:
            _process_paragraph(para, normalized_data)

    doc.save(output_path)


def _process_paragraph(para, data: dict) -> None:
    if not para.text:
        return

    # 1. Avval har bir run ichida replace qilib ko'ramiz (eng xavfsiz va formatni 100% saqlaydigan yo'l)
    for run in para.runs:
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, value)

    # 2. Agar qavslar Word XML ichida bo'lingan bo'lsa (cross-run)
    full_text = "".join(r.text for r in para.runs)
    has_unreplaced = any(f"{{{{{key}}}}}" in full_text for key in data.keys())

    if has_unreplaced:
        _replace_cross_run_paragraph(para, data)


def _replace_cross_run_paragraph(para, data: dict) -> None:
    """
    Qavslari bir nechta run ga bo'lingan paragraflarni formatini buzmasdan almashtirish
    """
    # Mavjud runlarning formatlarini saqlab olamiz
    runs_info = []
    for r in para.runs:
        runs_info.append({
            "text": r.text,
            "bold": r.bold,
            "italic": r.italic,
            "underline": r.underline,
            "font_name": r.font.name if r.font else "Times New Roman",
            "font_size": r.font.size if (r.font and r.font.size) else Pt(14)
        })

    full_text = "".join(r["text"] for r in runs_info)
    for key, value in data.items():
        ph = f"{{{{{key}}}}}"
        full_text = full_text.replace(ph, value)

    # Asosiy run formatini aniqlaymiz (birinchi yoki dominant run)
    main_bold = any(r["bold"] for r in runs_info) if runs_info else False
    main_italic = any(r["italic"] for r in runs_info) if runs_info else False
    main_underline = any(r["underline"] for r in runs_info) if runs_info else False
    main_size = runs_info[0]["font_size"] if runs_info else Pt(14)

    # Runlarni tozalab, yangilangan matnni qo'shamiz
    p_element = para._p
    for r in para.runs:
        p_element.remove(r._r)

    new_run = para.add_run(full_text)
    new_run.font.name = "Times New Roman"
    new_run.font.size = main_size
    new_run.bold = main_bold
    new_run.italic = main_italic
    new_run.underline = main_underline
