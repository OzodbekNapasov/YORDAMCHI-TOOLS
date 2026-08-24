"""
docx_generator.py — Test savollarini MS Word (.docx) formatiga o'tkazish moduli.
"""

import os
import io
import base64
import logging
from typing import List, Optional
from datetime import datetime
from PIL import Image

import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

try:
    from .xml_parser import Question
    from .pdf_generator import decode_image_base64, render_image_bytes
except ImportError:
    from xml_parser import Question
    from pdf_generator import decode_image_base64, render_image_bytes

logger = logging.getLogger(__name__)


def set_cell_background(cell, hex_color: str) -> None:
    """Jadval katagiga fon rang beradi."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def generate_docx(
    questions: List[Question],
    fan_name: str,
    with_answers: bool,
    output_path: str,
    test_id: Optional[str] = None
) -> str:
    """
    Savollarni yagona professional MS Word (.docx) hujjatiga yozadi.
    """
    doc = Document()
    
    # Hujjat hoshiyalari (margins)
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # ── Sarlavha ──
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(fan_name.upper())
    title_run.bold = True
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = RGBColor(30, 41, 59) # Slate 800

    # Meta ma'lumotlar
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_str = datetime.now().strftime("%d.%m.%Y %H:%M")
    variant_str = "Javoblar bilan" if with_answers else "Savollar (Javobsiz)"
    id_str = f"  |  Test ID: {test_id}" if test_id else ""
    meta_run = meta_p.add_run(f"Rejim: {variant_str}  |  Sana: {date_str}{id_str}")
    meta_run.font.name = 'Calibri'
    meta_run.font.size = Pt(9.5)
    meta_run.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph() # Bo'sh joy

    # ── Savollar ro'yxati ──
    for q in questions:
        # Savol matni
        qp = doc.add_paragraph()
        qp.paragraph_format.space_before = Pt(6)
        qp.paragraph_format.space_after = Pt(2)
        
        q_idx_run = qp.add_run(f"{q.index}. ")
        q_idx_run.bold = True
        q_idx_run.font.name = 'Calibri'
        q_idx_run.font.size = Pt(11)
        q_idx_run.font.color.rgb = RGBColor(15, 23, 42)

        q_txt_run = qp.add_run(q.question_text)
        q_txt_run.bold = True
        q_txt_run.font.name = 'Calibri'
        q_txt_run.font.size = Pt(11)
        q_txt_run.font.color.rgb = RGBColor(15, 23, 42)

        # Rasm mavjud bo'lsa (savoldan keyin)
        if q.image_base64:
            try:
                img_bytes = decode_image_base64(q.image_base64)
                pil_img = render_image_bytes(img_bytes)
                if pil_img:
                    temp_dir = os.path.dirname(output_path) or "."
                    temp_img_path = os.path.join(temp_dir, f"temp_docx_{q.index}.png")
                    pil_img.save(temp_img_path, format="PNG")

                    img_p = doc.add_paragraph()
                    img_p.paragraph_format.left_indent = Inches(0.3)
                    img_p.paragraph_format.space_before = Pt(2)
                    img_p.paragraph_format.space_after = Pt(4)
                    
                    run = img_p.add_run()
                    # Scale image to max width 4.5 inches
                    w_px, h_px = pil_img.size
                    w_in = min(4.5, w_px / 150.0)
                    run.add_picture(temp_img_path, width=Inches(w_in))

                    if os.path.exists(temp_img_path):
                        os.remove(temp_img_path)
            except Exception as e:
                logger.warning(f"Word ga rasm qo'shishda xatolik ({q.index}-savol): {e}")

        # Variantlar
        for opt_idx, v in enumerate(q.variants):
            letter = chr(ord("A") + opt_idx)
            vp = doc.add_paragraph()
            vp.paragraph_format.left_indent = Inches(0.3)
            vp.paragraph_format.space_before = Pt(1)
            vp.paragraph_format.space_after = Pt(1)

            v_letter_run = vp.add_run(f"{letter}) ")
            v_letter_run.bold = True
            v_letter_run.font.name = 'Calibri'
            v_letter_run.font.size = Pt(10.5)

            is_correct = with_answers and v.is_correct
            v_txt_run = vp.add_run(f"{v.text} *" if is_correct else v.text)
            v_txt_run.font.name = 'Calibri'
            v_txt_run.font.size = Pt(10.5)

            if is_correct:
                v_letter_run.font.color.rgb = RGBColor(16, 185, 129) # Emerald Green
                v_txt_run.bold = True
                v_txt_run.font.color.rgb = RGBColor(16, 185, 129)
            else:
                v_letter_run.font.color.rgb = RGBColor(71, 85, 105)
                v_txt_run.font.color.rgb = RGBColor(51, 65, 85)

    doc.save(output_path)
    return output_path


def generate_variants_docx(
    variants_questions: List[List[Question]],
    fan_name: str,
    with_answers: bool,
    output_path: str,
    test_id: Optional[str] = None
) -> str:
    """
    Bir nechta variantlarni bitta Word (.docx) fayliga jamlab yozadi.
    """
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    for var_idx, questions in enumerate(variants_questions):
        if var_idx > 0:
            doc.add_page_break()

        # Variant sarlavhasi
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run(f"{fan_name.upper()} - VARIANT {var_idx + 1}")
        title_run.bold = True
        title_run.font.name = 'Calibri'
        title_run.font.size = Pt(16)
        title_run.font.color.rgb = RGBColor(30, 41, 59)

        meta_p = doc.add_paragraph()
        meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_str = datetime.now().strftime("%d.%m.%Y %H:%M")
        variant_str = "Javoblar bilan" if with_answers else "Savollar (Javobsiz)"
        id_str = f"  |  Test ID: {test_id}" if test_id else ""
        meta_run = meta_p.add_run(f"Rejim: {variant_str}  |  Sana: {date_str}{id_str}")
        meta_run.font.name = 'Calibri'
        meta_run.font.size = Pt(9.5)
        meta_run.font.color.rgb = RGBColor(100, 116, 139)

        doc.add_paragraph()

        for q in questions:
            qp = doc.add_paragraph()
            qp.paragraph_format.space_before = Pt(6)
            qp.paragraph_format.space_after = Pt(2)

            q_idx_run = qp.add_run(f"{q.index}. ")
            q_idx_run.bold = True
            q_idx_run.font.name = 'Calibri'
            q_idx_run.font.size = Pt(11)

            q_txt_run = qp.add_run(q.question_text)
            q_txt_run.bold = True
            q_txt_run.font.name = 'Calibri'
            q_txt_run.font.size = Pt(11)

            if q.image_base64:
                try:
                    img_bytes = decode_image_base64(q.image_base64)
                    pil_img = render_image_bytes(img_bytes)
                    if pil_img:
                        temp_dir = os.path.dirname(output_path) or "."
                        temp_img_path = os.path.join(temp_dir, f"temp_var_{var_idx}_docx_{q.index}.png")
                        pil_img.save(temp_img_path, format="PNG")

                        img_p = doc.add_paragraph()
                        img_p.paragraph_format.left_indent = Inches(0.3)
                        run = img_p.add_run()
                        w_px, h_px = pil_img.size
                        w_in = min(4.5, w_px / 150.0)
                        run.add_picture(temp_img_path, width=Inches(w_in))

                        if os.path.exists(temp_img_path):
                            os.remove(temp_img_path)
                except Exception as e:
                    logger.warning(f"Word rasm xatosi ({q.index}-savol): {e}")

            for opt_idx, v in enumerate(q.variants):
                letter = chr(ord("A") + opt_idx)
                vp = doc.add_paragraph()
                vp.paragraph_format.left_indent = Inches(0.3)
                vp.paragraph_format.space_before = Pt(1)
                vp.paragraph_format.space_after = Pt(1)

                v_letter_run = vp.add_run(f"{letter}) ")
                v_letter_run.bold = True
                v_letter_run.font.name = 'Calibri'

                is_correct = with_answers and v.is_correct
                v_txt_run = vp.add_run(f"{v.text} *" if is_correct else v.text)
                v_txt_run.font.name = 'Calibri'

                if is_correct:
                    v_letter_run.font.color.rgb = RGBColor(16, 185, 129)
                    v_txt_run.bold = True
                    v_txt_run.font.color.rgb = RGBColor(16, 185, 129)
                else:
                    v_letter_run.font.color.rgb = RGBColor(71, 85, 105)

    doc.save(output_path)
    return output_path
